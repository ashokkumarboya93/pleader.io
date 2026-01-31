"""
Export utilities for Pleader AI
Supports PDF, DOCX, and TXT exports for chats and document analyses
"""

import logging
from typing import List, Dict, Any
from datetime import datetime
import io

# PDF generation
from reportlab.lib.pagesizes import letter, A4
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.units import inch
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, PageBreak
from reportlab.lib.enums import TA_LEFT, TA_CENTER

# DOCX generation
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

logger = logging.getLogger(__name__)


def export_chat_to_pdf(chat_data: Dict[str, Any]) -> bytes:
    """
    Export chat to PDF format
    
    Args:
        chat_data: Chat data including messages
        
    Returns:
        PDF file bytes
    """
    try:
        buffer = io.BytesIO()
        doc = SimpleDocTemplate(buffer, pagesize=letter)
        story = []
        styles = getSampleStyleSheet()
        
        # Custom styles
        title_style = ParagraphStyle(
            'CustomTitle',
            parent=styles['Heading1'],
            fontSize=18,
            textColor='#059669',  # Green theme
            spaceAfter=30,
            alignment=TA_CENTER
        )
        
        heading_style = ParagraphStyle(
            'CustomHeading',
            parent=styles['Heading2'],
            fontSize=12,
            textColor='#047857',
            spaceAfter=10,
            spaceBefore=15
        )
        
        message_style = ParagraphStyle(
            'MessageStyle',
            parent=styles['Normal'],
            fontSize=10,
            spaceAfter=12,
            leftIndent=20
        )
        
        # Title
        title = Paragraph("Pleader AI - Chat Export", title_style)
        story.append(title)
        story.append(Spacer(1, 0.2*inch))
        
        # Chat info
        chat_title = chat_data.get('title', 'Untitled Chat')
        created_at = chat_data.get('created_at', '')
        if isinstance(created_at, str):
            try:
                created_at = datetime.fromisoformat(created_at.replace('Z', '+00:00'))
                created_at = created_at.strftime('%B %d, %Y at %I:%M %p')
            except:
                pass
        
        info = Paragraph(f"<b>Chat:</b> {chat_title}<br/><b>Date:</b> {created_at}", styles['Normal'])
        story.append(info)
        story.append(Spacer(1, 0.3*inch))
        
        # Messages
        messages = chat_data.get('messages', [])
        for i, msg in enumerate(messages):
            sender = msg.get('sender', 'user')
            content = msg.get('content', '')
            timestamp = msg.get('timestamp', '')
            
            # Format timestamp
            if isinstance(timestamp, str) and timestamp:
                try:
                    ts = datetime.fromisoformat(timestamp.replace('Z', '+00:00'))
                    timestamp = ts.strftime('%I:%M %p')
                except:
                    timestamp = ''
            
            # Sender heading
            sender_label = "You" if sender == "user" else "Pleader AI"
            heading_text = f"{sender_label} {f'({timestamp})' if timestamp else ''}"
            story.append(Paragraph(heading_text, heading_style))
            
            # Message content
            content_safe = content.replace('<', '&lt;').replace('>', '&gt;').replace('\n', '<br/>')
            story.append(Paragraph(content_safe, message_style))
            story.append(Spacer(1, 0.1*inch))
        
        # Build PDF
        doc.build(story)
        buffer.seek(0)
        
        logger.info(f"Exported chat to PDF with {len(messages)} messages")
        return buffer.getvalue()
        
    except Exception as e:
        logger.error(f"Error exporting chat to PDF: {e}")
        raise Exception(f"Failed to export chat to PDF: {str(e)}")


def export_chat_to_docx(chat_data: Dict[str, Any]) -> bytes:
    """
    Export chat to DOCX format
    
    Args:
        chat_data: Chat data including messages
        
    Returns:
        DOCX file bytes
    """
    try:
        doc = Document()
        
        # Title
        title = doc.add_heading('Pleader AI - Chat Export', level=0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        title_run = title.runs[0]
        title_run.font.color.rgb = RGBColor(5, 150, 105)  # Green theme
        
        # Chat info
        chat_title = chat_data.get('title', 'Untitled Chat')
        created_at = chat_data.get('created_at', '')
        if isinstance(created_at, str):
            try:
                created_at = datetime.fromisoformat(created_at.replace('Z', '+00:00'))
                created_at = created_at.strftime('%B %d, %Y at %I:%M %p')
            except:
                pass
        
        info_para = doc.add_paragraph()
        info_para.add_run('Chat: ').bold = True
        info_para.add_run(f"{chat_title}\n")
        info_para.add_run('Date: ').bold = True
        info_para.add_run(str(created_at))
        
        doc.add_paragraph()  # Spacer
        
        # Messages
        messages = chat_data.get('messages', [])
        for msg in messages:
            sender = msg.get('sender', 'user')
            content = msg.get('content', '')
            timestamp = msg.get('timestamp', '')
            
            # Format timestamp
            if isinstance(timestamp, str) and timestamp:
                try:
                    ts = datetime.fromisoformat(timestamp.replace('Z', '+00:00'))
                    timestamp = ts.strftime('%I:%M %p')
                except:
                    timestamp = ''
            
            # Sender heading
            sender_label = "You" if sender == "user" else "Pleader AI"
            heading_text = f"{sender_label} {f'({timestamp})' if timestamp else ''}"
            heading = doc.add_heading(heading_text, level=2)
            heading_run = heading.runs[0]
            heading_run.font.color.rgb = RGBColor(4, 120, 87)
            
            # Message content
            content_para = doc.add_paragraph(content)
            content_para.paragraph_format.left_indent = Inches(0.3)
            
            doc.add_paragraph()  # Spacer
        
        # Save to buffer
        buffer = io.BytesIO()
        doc.save(buffer)
        buffer.seek(0)
        
        logger.info(f"Exported chat to DOCX with {len(messages)} messages")
        return buffer.getvalue()
        
    except Exception as e:
        logger.error(f"Error exporting chat to DOCX: {e}")
        raise Exception(f"Failed to export chat to DOCX: {str(e)}")


def export_chat_to_txt(chat_data: Dict[str, Any]) -> str:
    """
    Export chat to plain text format
    
    Args:
        chat_data: Chat data including messages
        
    Returns:
        Plain text string
    """
    try:
        lines = []
        lines.append("=" * 60)
        lines.append("PLEADER AI - CHAT EXPORT")
        lines.append("=" * 60)
        lines.append("")
        
        # Chat info
        chat_title = chat_data.get('title', 'Untitled Chat')
        created_at = chat_data.get('created_at', '')
        if isinstance(created_at, str):
            try:
                created_at = datetime.fromisoformat(created_at.replace('Z', '+00:00'))
                created_at = created_at.strftime('%B %d, %Y at %I:%M %p')
            except:
                pass
        
        lines.append(f"Chat: {chat_title}")
        lines.append(f"Date: {created_at}")
        lines.append("")
        lines.append("-" * 60)
        lines.append("")
        
        # Messages
        messages = chat_data.get('messages', [])
        for msg in messages:
            sender = msg.get('sender', 'user')
            content = msg.get('content', '')
            timestamp = msg.get('timestamp', '')
            
            # Format timestamp
            if isinstance(timestamp, str) and timestamp:
                try:
                    ts = datetime.fromisoformat(timestamp.replace('Z', '+00:00'))
                    timestamp = ts.strftime('%I:%M %p')
                except:
                    timestamp = ''
            
            # Sender heading
            sender_label = "YOU" if sender == "user" else "PLEADER AI"
            heading = f"{sender_label} {f'({timestamp})' if timestamp else ''}"
            lines.append(heading)
            lines.append("-" * len(heading))
            lines.append(content)
            lines.append("")
            lines.append("")
        
        text = "\n".join(lines)
        logger.info(f"Exported chat to TXT with {len(messages)} messages")
        return text
        
    except Exception as e:
        logger.error(f"Error exporting chat to TXT: {e}")
        raise Exception(f"Failed to export chat to TXT: {str(e)}")


def export_analysis_to_pdf(analysis_data: Dict[str, Any]) -> bytes:
    """
    Export document analysis to PDF format
    
    Args:
        analysis_data: Document analysis data
        
    Returns:
        PDF file bytes
    """
    try:
        buffer = io.BytesIO()
        doc = SimpleDocTemplate(buffer, pagesize=letter)
        story = []
        styles = getSampleStyleSheet()
        
        # Custom styles
        title_style = ParagraphStyle(
            'CustomTitle',
            parent=styles['Heading1'],
            fontSize=18,
            textColor='#059669',
            spaceAfter=30,
            alignment=TA_CENTER
        )
        
        section_style = ParagraphStyle(
            'SectionStyle',
            parent=styles['Heading2'],
            fontSize=14,
            textColor='#047857',
            spaceAfter=15,
            spaceBefore=20
        )
        
        # Title
        title = Paragraph("Pleader AI - Document Analysis", title_style)
        story.append(title)
        story.append(Spacer(1, 0.2*inch))
        
        # Document info
        filename = analysis_data.get('filename', 'Unknown')
        uploaded_at = analysis_data.get('uploaded_at', '')
        if isinstance(uploaded_at, str):
            try:
                uploaded_at = datetime.fromisoformat(uploaded_at.replace('Z', '+00:00'))
                uploaded_at = uploaded_at.strftime('%B %d, %Y at %I:%M %p')
            except:
                pass
        
        info = Paragraph(f"<b>Document:</b> {filename}<br/><b>Analyzed:</b> {uploaded_at}", styles['Normal'])
        story.append(info)
        story.append(Spacer(1, 0.3*inch))
        
        # Analysis result
        analysis = analysis_data.get('analysis_result', {})
        
        # Full analysis
        full_analysis = analysis.get('full_analysis', 'No analysis available')
        story.append(Paragraph("Analysis", section_style))
        
        # Split into paragraphs
        paragraphs = full_analysis.split('\n\n')
        for para in paragraphs:
            if para.strip():
                para_safe = para.replace('<', '&lt;').replace('>', '&gt;').replace('\n', '<br/>')
                story.append(Paragraph(para_safe, styles['Normal']))
                story.append(Spacer(1, 0.1*inch))
        
        # Build PDF
        doc.build(story)
        buffer.seek(0)
        
        logger.info(f"Exported analysis to PDF for {filename}")
        return buffer.getvalue()
        
    except Exception as e:
        logger.error(f"Error exporting analysis to PDF: {e}")
        raise Exception(f"Failed to export analysis to PDF: {str(e)}")


def export_analysis_to_docx(analysis_data: Dict[str, Any]) -> bytes:
    """
    Export document analysis to DOCX format
    
    Args:
        analysis_data: Document analysis data
        
    Returns:
        DOCX file bytes
    """
    try:
        doc = Document()
        
        # Title
        title = doc.add_heading('Pleader AI - Document Analysis', level=0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        title_run = title.runs[0]
        title_run.font.color.rgb = RGBColor(5, 150, 105)
        
        # Document info
        filename = analysis_data.get('filename', 'Unknown')
        uploaded_at = analysis_data.get('uploaded_at', '')
        if isinstance(uploaded_at, str):
            try:
                uploaded_at = datetime.fromisoformat(uploaded_at.replace('Z', '+00:00'))
                uploaded_at = uploaded_at.strftime('%B %d, %Y at %I:%M %p')
            except:
                pass
        
        info_para = doc.add_paragraph()
        info_para.add_run('Document: ').bold = True
        info_para.add_run(f"{filename}\n")
        info_para.add_run('Analyzed: ').bold = True
        info_para.add_run(str(uploaded_at))
        
        doc.add_paragraph()
        
        # Analysis result
        analysis = analysis_data.get('analysis_result', {})
        full_analysis = analysis.get('full_analysis', 'No analysis available')
        
        doc.add_heading('Analysis', level=1)
        
        # Add analysis text
        paragraphs = full_analysis.split('\n\n')
        for para in paragraphs:
            if para.strip():
                doc.add_paragraph(para.strip())
        
        # Save to buffer
        buffer = io.BytesIO()
        doc.save(buffer)
        buffer.seek(0)
        
        logger.info(f"Exported analysis to DOCX for {filename}")
        return buffer.getvalue()
        
    except Exception as e:
        logger.error(f"Error exporting analysis to DOCX: {e}")
        raise Exception(f"Failed to export analysis to DOCX: {str(e)}")


def export_analysis_to_txt(analysis_data: Dict[str, Any]) -> str:
    """
    Export document analysis to plain text format
    
    Args:
        analysis_data: Document analysis data
        
    Returns:
        Plain text string
    """
    try:
        lines = []
        lines.append("=" * 60)
        lines.append("PLEADER AI - DOCUMENT ANALYSIS")
        lines.append("=" * 60)
        lines.append("")
        
        # Document info
        filename = analysis_data.get('filename', 'Unknown')
        uploaded_at = analysis_data.get('uploaded_at', '')
        if isinstance(uploaded_at, str):
            try:
                uploaded_at = datetime.fromisoformat(uploaded_at.replace('Z', '+00:00'))
                uploaded_at = uploaded_at.strftime('%B %d, %Y at %I:%M %p')
            except:
                pass
        
        lines.append(f"Document: {filename}")
        lines.append(f"Analyzed: {uploaded_at}")
        lines.append("")
        lines.append("-" * 60)
        lines.append("")
        
        # Analysis
        analysis = analysis_data.get('analysis_result', {})
        full_analysis = analysis.get('full_analysis', 'No analysis available')
        
        lines.append("ANALYSIS")
        lines.append("-" * 60)
        lines.append("")
        lines.append(full_analysis)
        lines.append("")
        
        text = "\n".join(lines)
        logger.info(f"Exported analysis to TXT for {filename}")
        return text
        
    except Exception as e:
        logger.error(f"Error exporting analysis to TXT: {e}")
        raise Exception(f"Failed to export analysis to TXT: {str(e)}")
