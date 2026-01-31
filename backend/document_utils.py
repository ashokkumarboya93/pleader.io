"""
Document extraction utilities for Pleader AI
Supports PDF, DOCX, TXT, and images (JPG, PNG) with OCR
"""

import logging
from typing import Optional
from pathlib import Path
import io

# PDF extraction
from pypdf import PdfReader

# DOCX extraction
from docx import Document

# Image OCR
from PIL import Image
import pytesseract

logger = logging.getLogger(__name__)


def extract_text_from_pdf(file_content: bytes) -> str:
    """
    Extract text from PDF file
    
    Args:
        file_content: PDF file bytes
        
    Returns:
        Extracted text
    """
    try:
        pdf_file = io.BytesIO(file_content)
        reader = PdfReader(pdf_file)
        
        text_parts = []
        for page_num, page in enumerate(reader.pages):
            try:
                text = page.extract_text()
                if text:
                    text_parts.append(f"[Page {page_num + 1}]\n{text}")
            except Exception as e:
                logger.warning(f"Error extracting page {page_num + 1}: {e}")
        
        extracted_text = "\n\n".join(text_parts)
        logger.info(f"Extracted {len(extracted_text)} characters from PDF ({len(reader.pages)} pages)")
        
        return extracted_text
        
    except Exception as e:
        logger.error(f"Error extracting PDF: {e}")
        return f"Error extracting PDF: {str(e)}"


def extract_text_from_docx(file_content: bytes) -> str:
    """
    Extract text from DOCX file
    
    Args:
        file_content: DOCX file bytes
        
    Returns:
        Extracted text
    """
    try:
        docx_file = io.BytesIO(file_content)
        doc = Document(docx_file)
        
        text_parts = []
        
        # Extract paragraphs
        for para in doc.paragraphs:
            if para.text.strip():
                text_parts.append(para.text)
        
        # Extract tables
        for table in doc.tables:
            for row in table.rows:
                row_text = " | ".join(cell.text.strip() for cell in row.cells)
                if row_text.strip():
                    text_parts.append(row_text)
        
        extracted_text = "\n\n".join(text_parts)
        logger.info(f"Extracted {len(extracted_text)} characters from DOCX")
        
        return extracted_text
        
    except Exception as e:
        logger.error(f"Error extracting DOCX: {e}")
        return f"Error extracting DOCX: {str(e)}"


def extract_text_from_image(file_content: bytes, filename: str = "") -> str:
    """
    Extract text from image using OCR (JPG, PNG, etc.)
    
    Args:
        file_content: Image file bytes
        filename: Original filename (for logging)
        
    Returns:
        Extracted text via OCR
    """
    try:
        image = Image.open(io.BytesIO(file_content))
        
        # Convert to RGB if needed
        if image.mode != 'RGB':
            image = image.convert('RGB')
        
        # Perform OCR
        text = pytesseract.image_to_string(image)
        
        logger.info(f"Extracted {len(text)} characters from image {filename} using OCR")
        
        if not text.strip():
            return "No text could be extracted from this image. The image may not contain readable text."
        
        return text
        
    except Exception as e:
        logger.error(f"Error extracting text from image: {e}")
        return f"Error performing OCR on image: {str(e)}"


def extract_text_from_txt(file_content: bytes) -> str:
    """
    Extract text from TXT file
    
    Args:
        file_content: Text file bytes
        
    Returns:
        Decoded text
    """
    try:
        # Try UTF-8 first
        try:
            text = file_content.decode('utf-8')
        except UnicodeDecodeError:
            # Fallback to latin-1
            text = file_content.decode('latin-1', errors='ignore')
        
        logger.info(f"Extracted {len(text)} characters from TXT file")
        return text
        
    except Exception as e:
        logger.error(f"Error extracting TXT: {e}")
        return f"Error extracting text: {str(e)}"


def extract_text_from_file(file_content: bytes, filename: str) -> str:
    """
    Extract text from any supported file type
    
    Args:
        file_content: File bytes
        filename: Original filename
        
    Returns:
        Extracted text
    """
    file_ext = filename.lower().split('.')[-1] if '.' in filename else ''
    
    if file_ext == 'pdf':
        return extract_text_from_pdf(file_content)
    elif file_ext in ['docx', 'doc']:
        if file_ext == 'doc':
            logger.warning("Legacy .doc format detected. Only .docx is fully supported.")
            return "Legacy .doc format is not supported. Please convert to .docx format."
        return extract_text_from_docx(file_content)
    elif file_ext in ['txt', 'text']:
        return extract_text_from_txt(file_content)
    elif file_ext in ['jpg', 'jpeg', 'png', 'bmp', 'tiff']:
        return extract_text_from_image(file_content, filename)
    else:
        logger.warning(f"Unsupported file type: {file_ext}")
        return f"Unsupported file type: {file_ext}. Supported types: PDF, DOCX, TXT, JPG, PNG"


def validate_file_type(filename: str) -> bool:
    """
    Check if file type is supported
    
    Args:
        filename: File name
        
    Returns:
        True if supported, False otherwise
    """
    supported_extensions = ['pdf', 'docx', 'doc', 'txt', 'text', 'jpg', 'jpeg', 'png', 'bmp', 'tiff']
    file_ext = filename.lower().split('.')[-1] if '.' in filename else ''
    return file_ext in supported_extensions
